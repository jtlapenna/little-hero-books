#!/usr/bin/env python3
"""Extract Semrush SEO Content Template DOCX files into text and metadata.

The script is capture-only. It preserves the original DOCX files and writes
searchable text plus package metadata for future analysis after Semrush access
expires.
"""

from __future__ import annotations

import csv
import datetime as dt
import hashlib
import json
import zipfile
from pathlib import Path
from typing import Any

from docx import Document


ARCHIVE_DIR = Path(__file__).resolve().parents[1]
ORIGINALS_DIR = ARCHIVE_DIR / "originals"
TEXT_DIR = ARCHIVE_DIR / "extracted-text"
METADATA_DIR = ARCHIVE_DIR / "metadata"


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def core_properties(document: Document) -> dict[str, Any]:
    props = document.core_properties
    result: dict[str, Any] = {}
    for name in [
        "author",
        "category",
        "comments",
        "content_status",
        "created",
        "identifier",
        "keywords",
        "language",
        "last_modified_by",
        "last_printed",
        "modified",
        "revision",
        "subject",
        "title",
        "version",
    ]:
        value = getattr(props, name)
        if isinstance(value, (dt.datetime, dt.date, dt.time)):
            value = value.isoformat()
        result[name] = value
    return result


def iter_block_text(document: Document) -> tuple[list[str], list[dict[str, Any]]]:
    lines: list[str] = []
    tables: list[dict[str, Any]] = []

    for paragraph in document.paragraphs:
        text = paragraph.text
        if text.strip():
            lines.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        rows: list[list[str]] = []
        lines.append("")
        lines.append(f"[Table {table_index}]")
        for row in table.rows:
            values = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            rows.append(values)
            lines.append(" | ".join(values))
        tables.append(
            {
                "table_index": table_index,
                "row_count": len(rows),
                "column_count": max((len(row) for row in rows), default=0),
                "rows": rows,
            }
        )

    return lines, tables


def package_parts(path: Path) -> list[dict[str, Any]]:
    parts: list[dict[str, Any]] = []
    with zipfile.ZipFile(path) as archive:
        for info in sorted(archive.infolist(), key=lambda item: item.filename):
            parts.append(
                {
                    "filename": info.filename,
                    "file_size": info.file_size,
                    "compress_size": info.compress_size,
                    "crc": f"{info.CRC:08x}",
                }
            )
    return parts


def extract_docx(path: Path) -> dict[str, Any]:
    document = Document(path)
    lines, tables = iter_block_text(document)

    text_path = TEXT_DIR / f"{path.stem}.txt"
    text_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")

    parts = package_parts(path)
    part_manifest_path = METADATA_DIR / f"{path.stem}.package-parts.json"
    part_manifest_path.write_text(
        json.dumps(parts, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )

    return {
        "file_name": path.name,
        "relative_path": str(path.relative_to(ARCHIVE_DIR)),
        "size_bytes": path.stat().st_size,
        "sha256": sha256_file(path),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "extracted_line_count": len(lines),
        "tables": tables,
        "core_properties": core_properties(document),
        "extracted_text_relative_path": str(text_path.relative_to(ARCHIVE_DIR)),
        "extracted_text_sha256": sha256_file(text_path),
        "package_part_count": len(parts),
        "package_parts_relative_path": str(part_manifest_path.relative_to(ARCHIVE_DIR)),
        "package_parts_sha256": sha256_file(part_manifest_path),
    }


def write_csv(workbooks: list[dict[str, Any]]) -> None:
    csv_path = METADATA_DIR / "content_templates_manifest.csv"
    with csv_path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(
            handle,
            fieldnames=[
                "file_name",
                "relative_path",
                "size_bytes",
                "sha256",
                "paragraph_count",
                "table_count",
                "extracted_line_count",
                "extracted_text_relative_path",
                "package_part_count",
                "package_parts_relative_path",
                "title",
                "subject",
                "keywords",
                "created",
                "modified",
            ],
        )
        writer.writeheader()
        for workbook in workbooks:
            props = workbook["core_properties"]
            writer.writerow(
                {
                    "file_name": workbook["file_name"],
                    "relative_path": workbook["relative_path"],
                    "size_bytes": workbook["size_bytes"],
                    "sha256": workbook["sha256"],
                    "paragraph_count": workbook["paragraph_count"],
                    "table_count": workbook["table_count"],
                    "extracted_line_count": workbook["extracted_line_count"],
                    "extracted_text_relative_path": workbook["extracted_text_relative_path"],
                    "package_part_count": workbook["package_part_count"],
                    "package_parts_relative_path": workbook["package_parts_relative_path"],
                    "title": props["title"],
                    "subject": props["subject"],
                    "keywords": props["keywords"],
                    "created": props["created"],
                    "modified": props["modified"],
                }
            )


def write_checksums(manifest: dict[str, Any]) -> None:
    paths: list[Path] = []
    for item in manifest["content_templates"]:
        paths.append(ARCHIVE_DIR / item["relative_path"])
        paths.append(ARCHIVE_DIR / item["extracted_text_relative_path"])
        paths.append(ARCHIVE_DIR / item["package_parts_relative_path"])

    checksums_path = METADATA_DIR / "docx_checksums.sha256"
    with checksums_path.open("w", encoding="utf-8") as handle:
        for path in sorted(paths):
            handle.write(f"{sha256_file(path)}  {path.relative_to(ARCHIVE_DIR)}\n")


def main() -> None:
    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    METADATA_DIR.mkdir(parents=True, exist_ok=True)

    content_templates = [extract_docx(path) for path in sorted(ORIGINALS_DIR.glob("*.docx"))]
    manifest: dict[str, Any] = {
        "capture_name": "Semrush Position Tracking and SEO Content Template capture",
        "archive_date": "2026-06-30",
        "domain": "littleherolabs.com",
        "content_templates": content_templates,
    }

    manifest_path = METADATA_DIR / "content_templates_manifest.json"
    manifest_path.write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    write_csv(content_templates)
    write_checksums(manifest)


if __name__ == "__main__":
    main()
